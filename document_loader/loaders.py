"""
通用文档加载器：支持 PDF、DOCX、Markdown、TXT、HTML、CSV
所有函数返回 List[Document]，统一接口供后续分块使用。
"""

import csv
from pathlib import Path
from typing import List, Dict, Any
import logging

import pypdf
import docx
import markdown
from bs4 import BeautifulSoup

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ---------- 统一数据结构 ----------
class Document:
    """RAG 文档最小单元，page_content 存储文本，metadata 携带来源信息"""
    __slots__ = ("page_content", "metadata")

    def __init__(self, page_content: str, metadata: Dict[str, Any]):
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        src = self.metadata.get("source", "unknown")
        return f"Document({src}, len={len(self.page_content)})"


# ---------- 各格式解析器 ----------
def load_pdf(file_path: str) -> List[Document]:
    """提取 PDF 每页文本作为一个 Document"""
    docs = []
    try:
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        page_content=text.strip(),
                        metadata={
                            "source": file_path,
                            "page": page_num,
                            "type": "pdf"
                        }
                    ))
    except Exception as e:
        logger.error(f"解析 PDF 失败: {file_path} - {e}")
    return docs


def load_docx(file_path: str) -> List[Document]:
    """提取 Word 文档所有段落（合并为一个 Document）"""
    try:
        doc = docx.Document(file_path)
        full_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
        if full_text.strip():
            return [Document(
                page_content=full_text.strip(),
                metadata={"source": file_path, "type": "docx"}
            )]
    except Exception as e:
        logger.error(f"解析 DOCX 失败: {file_path} - {e}")
    return []


def load_md(file_path: str) -> List[Document]:
    """将 Markdown 转为纯文本（去除格式标记）"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, "lxml")
        plain_text = soup.get_text(separator="\n")
        if plain_text.strip():
            return [Document(
                page_content=plain_text.strip(),
                metadata={"source": file_path, "type": "markdown"}
            )]
    except Exception as e:
        logger.error(f"解析 MD 失败: {file_path} - {e}")
    return []


def load_txt(file_path: str) -> List[Document]:
    """直接读取文本文件，自动尝试多种编码"""
    # 常用编码尝试列表
    encodings = ["utf-8", "gbk", "gb2312", "latin-1", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read()
            if text.strip():
                return [Document(
                    page_content=text.strip(),
                    metadata={"source": file_path, "type": "txt", "encoding": enc}
                )]
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"读取 TXT 时出错: {file_path} - {e}")
            return []
    logger.error(f"无法解码文件: {file_path}")
    return []


def load_html(file_path: str) -> List[Document]:
    """从 HTML 文件中提取纯文本"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "lxml")
        text = soup.get_text(separator="\n")
        if text.strip():
            return [Document(
                page_content=text.strip(),
                metadata={"source": file_path, "type": "html"}
            )]
    except Exception as e:
        logger.error(f"解析 HTML 失败: {file_path} - {e}")
    return []


def load_csv(file_path: str, delimiter: str = ",") -> List[Document]:
    """
    将 CSV 每一行转换为一个 Document，内容为 '列名: 值' 对。
    可通过 delimiter 指定分隔符（如 '\t' 用于 TSV）。
    """
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            # 自动检测是否有表头
            sample = f.read(1024)
            f.seek(0)
            has_header = csv.Sniffer().has_header(sample)
            if has_header:
                reader = csv.DictReader(f, delimiter=delimiter)
                for i, row in enumerate(reader):
                    content = "\n".join(
                        f"{k}: {v}" for k, v in row.items() if v
                    )
                    if content.strip():
                        docs.append(Document(
                            page_content=content,
                            metadata={"source": file_path, "row": i, "type": "csv"}
                        ))
            else:
                reader = csv.reader(f, delimiter=delimiter)
                for i, row in enumerate(reader):
                    content = ", ".join(row)
                    if content.strip():
                        docs.append(Document(
                            page_content=content,
                            metadata={"source": file_path, "row": i, "type": "csv"}
                        ))
    except Exception as e:
        logger.error(f"解析 CSV 失败: {file_path} - {e}")
    return docs


# ---------- 扩展名 -> 解析器 映射表 ----------
LOADER_MAP: Dict[str, Any] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,          # 旧 .doc 可能失败，建议转为 docx
    ".md": load_md,
    ".markdown": load_md,
    ".txt": load_txt,
    ".html": load_html,
    ".htm": load_html,
    ".csv": load_csv,
    ".tsv": lambda fp: load_csv(fp, delimiter="\t"),   # TSV 变体
}


# ---------- 公共接口 ----------
def load_document(file_path: str) -> List[Document]:
    """根据文件扩展名自动调用对应的解析器"""
    ext = Path(file_path).suffix.lower()
    loader = LOADER_MAP.get(ext)
    if not loader:
        logger.warning(f"不支持的文件类型: {ext} （文件: {file_path}）")
        return []
    return loader(file_path)


def load_documents(directory: str, recursive: bool = True) -> List[Document]:
    """
    遍历文件夹，加载所有支持的文档。
    recursive=True 会搜索子文件夹。
    """
    all_docs = []
    base = Path(directory)
    if not base.exists():
        logger.error(f"目录不存在: {directory}")
        return all_docs

    pattern = "**/*" if recursive else "*"
    for file_path in base.glob(pattern):
        if file_path.is_file():
            all_docs.extend(load_document(str(file_path)))
    logger.info(f"共加载 {len(all_docs)} 个文档块")
    return all_docs